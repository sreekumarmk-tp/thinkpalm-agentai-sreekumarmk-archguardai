import os
import re
import requests
import io
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

def clean_for_pdf(text: str) -> str:
    """Removes or replaces characters that fpdf2 might not handle well with default fonts."""
    # Replace common markdown bold/italic/code symbols for a cleaner text-only look in simple PDF
    text = text.replace('**', '').replace('*', '').replace('`', '')
    # Remove characters outside the Latin-1 range as FPDF's default fonts are limited
    return "".join(i if ord(i) < 256 else " " for i in text)

from src.utils.mermaid_cleanup import cleanup_mermaid_code, get_mermaid_url

def get_mermaid_image_path(mermaid_code: str, run_id: str, index: int) -> str:
    """Fetches a Mermaid diagram image from mermaid.ink and saves it locally."""
    
    # Apply standardized cleanup
    cleaned = cleanup_mermaid_code(mermaid_code)
    url = get_mermaid_url(cleaned)
    
    try:
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            path = f"/tmp/mermaid_{run_id}_{index}.png"
            with open(path, "wb") as f:
                f.write(response.content)
            return path
    except Exception:
        pass
    return None

def export_to_docx(markdown_content: str, run_id: str) -> bytes:
    """Converts markdown content to a Word Document as bytes."""
    doc = Document()
    doc.add_heading('ArchGuard AI: Architecture Review Report', 0)
    
    blocks = re.split(r'(\n\n)', markdown_content)
    mermaid_idx = 0

    for block in blocks:
        block = block.strip()
        if not block:
            continue
            
        if block.startswith('```mermaid'):
            code = re.sub(r'^```mermaid\s*\n|```$', '', block, flags=re.MULTILINE).strip()
            img_path = get_mermaid_image_path(code, run_id, mermaid_idx)
            if img_path:
                doc.add_picture(img_path, width=Inches(5.0))
                if os.path.exists(img_path):
                    os.remove(img_path)
            else:
                doc.add_paragraph("[Diagram: Mermaid source code below]")
                doc.add_paragraph(code)
            mermaid_idx += 1
            continue

        if '|' in block and '-' in block:
            lines = block.split('\n')
            if len(lines) >= 2 and '|' in lines[0] and '|' in lines[1]:
                rows = []
                for line in lines:
                    if line.strip().startswith('|'):
                        cells = [c.strip() for c in line.split('|') if c.strip()]
                        if cells and not all(c.strip('-') == '' for c in cells):
                            rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(rows):
                        for j, cell_data in enumerate(row_data):
                            if j < len(table.columns):
                                table.cell(i, j).text = cell_data
                    continue

        lines = block.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                clean_line = line.replace('**', '').replace('*', '').replace('`', '')
                doc.add_paragraph(clean_line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def export_to_pdf(markdown_content: str, run_id: str) -> bytes:
    """Converts markdown content to a PDF as bytes."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 15, txt="ArchGuard AI: Architecture Review Report", ln=True, align='C')
    pdf.ln(5)
    
    mermaid_idx = 0
    blocks = re.split(r'(\n\n)', markdown_content)
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        if block.startswith('```mermaid'):
            code = re.sub(r'^```mermaid\s*\n|```$', '', block, flags=re.MULTILINE).strip()
            img_path = get_mermaid_image_path(code, run_id, mermaid_idx)
            if img_path:
                pdf.image(img_path, w=150)
                pdf.ln(2)
                if os.path.exists(img_path):
                    os.remove(img_path)
            else:
                pdf.set_font("Helvetica", "I", 10)
                pdf.multi_cell(0, 7, txt="[Diagram: Mermaid code follows]")
                pdf.set_font("Courier", size=8)
                pdf.multi_cell(0, 5, txt=code)
            mermaid_idx += 1
            continue

        if '|' in block and '-' in block:
            lines = block.split('\n')
            if len(lines) >= 2 and '|' in lines[0] and '|' in lines[1]:
                pdf.set_font("Helvetica", "B", 10)
                rows = []
                for line in lines:
                    if line.strip().startswith('|'):
                        cells = [c.strip() for c in line.split('|') if c.strip()]
                        if cells and not all(c.strip('-') == '' for c in cells):
                            rows.append(cells)
                
                if rows:
                    col_width = 190 / len(rows[0])
                    for i, row in enumerate(rows):
                        if i == 0: pdf.set_font("Helvetica", "B", 10)
                        else: pdf.set_font("Helvetica", size=9)
                        for cell in row:
                            pdf.cell(col_width, 10, txt=clean_for_pdf(cell), border=1)
                        pdf.ln()
                    continue

        lines = block.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith('# '):
                pdf.set_font("Helvetica", "B", 16)
                pdf.multi_cell(0, 10, txt=clean_for_pdf(line[2:]))
            elif line.startswith('## '):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 9, txt=clean_for_pdf(line[3:]))
            elif line.startswith('### '):
                pdf.set_font("Helvetica", "B", 12)
                pdf.multi_cell(0, 8, txt=clean_for_pdf(line[4:]))
            else:
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 7, txt=clean_for_pdf(line))
            
    return pdf.output()
